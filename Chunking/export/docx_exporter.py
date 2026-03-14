from __future__ import annotations

import re
from pathlib import Path
from typing import List

from docx import Document

from Chunking.chunking.models import Chunk, DocumentMetadata


# ============================================================================
# XML 1.0 valid character sanitizer
# ============================================================================
#
# python-docx writes XML under the hood. Some PDF extractions may contain
# invisible control characters that are valid in Python strings but invalid
# in XML / DOCX documents.
#
# This sanitizer removes those characters before writing text to Word.
# ============================================================================

INVALID_XML_CHARS_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")


def sanitize_for_docx(text: str) -> str:
    """
    Remove characters that are not valid in XML 1.0.

    Why this is necessary
    ---------------------
    - PDF extraction sometimes introduces hidden control characters
    - python-docx serializes content as XML
    - XML rejects NULL bytes and several control characters

    We intentionally preserve normal whitespace such as:
    - newline
    - carriage return
    - tab

    Parameters
    ----------
    text : str
        Input text that may contain invalid XML characters.

    Returns
    -------
    str
        Sanitized text safe to write into a DOCX document.
    """
    if not text:
        return ""

    return INVALID_XML_CHARS_RE.sub("", text)


class DocxInspectionExporter:
    """
    Export chunk inspection results into a DOCX file.

    Why this exporter matters
    -------------------------
    Human inspection remains extremely important for validating:
    - chunk boundaries
    - text cleanliness
    - metadata richness
    - structural traceability
    - semantic coherence

    This exporter is intentionally inspection-oriented rather than
    presentation-oriented.
    """

    def write_chunks_docx(
        self,
        document_metadata: DocumentMetadata,
        strategy_name: str,
        chunks: List[Chunk],
        output_path: Path,
    ) -> None:
        """
        Write a chunk inspection DOCX file.

        Report structure
        ----------------
        The report includes:
        - document-level summary
        - one section per chunk
        - explicit structural fields
        - metadata
        - visible chunk text
        - optional embedding text when different from visible text

        Parameters
        ----------
        document_metadata : DocumentMetadata
            High-level source document metadata.

        strategy_name : str
            Name of the chunking strategy used.

        chunks : List[Chunk]
            Final chunk list in document order.

        output_path : Path
            Destination DOCX path.
        """
        document = Document()

        self._write_document_header(
            document=document,
            document_metadata=document_metadata,
            strategy_name=strategy_name,
            chunks=chunks,
        )

        for index, chunk in enumerate(chunks, start=1):
            self._write_chunk_section(
                document=document,
                chunk=chunk,
                index=index,
            )

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)

    def _write_document_header(
        self,
        document: Document,
        document_metadata: DocumentMetadata,
        strategy_name: str,
        chunks: List[Chunk],
    ) -> None:
        """
        Write the document-level inspection summary.

        Parameters
        ----------
        document : Document
            Target DOCX document object.

        document_metadata : DocumentMetadata
            High-level source document metadata.

        strategy_name : str
            Name of the chunking strategy used.

        chunks : List[Chunk]
            Final chunk list.
        """
        document.add_heading("Chunk Inspection Report", level=1)
        document.add_paragraph(f"Document: {sanitize_for_docx(document_metadata.title)}")
        document.add_paragraph(f"File: {sanitize_for_docx(document_metadata.file_name)}")
        document.add_paragraph(f"Strategy: {sanitize_for_docx(strategy_name)}")
        document.add_paragraph(f"Total chunks: {len(chunks)}")

        char_counts = [
            getattr(chunk, "char_count", len(chunk.text or ""))
            for chunk in chunks
        ]

        if char_counts:
            document.add_paragraph(f"Minimum characters: {min(char_counts)}")
            document.add_paragraph(f"Maximum characters: {max(char_counts)}")
            document.add_paragraph(
                f"Average characters: {sum(char_counts) / len(char_counts):.2f}"
            )

    def _write_chunk_section(
        self,
        document: Document,
        chunk: Chunk,
        index: int,
    ) -> None:
        """
        Write one chunk section into the inspection DOCX.

        Parameters
        ----------
        document : Document
            Target DOCX document object.

        chunk : Chunk
            Chunk to render.

        index : int
            Human-friendly chunk number in the report.
        """
        document.add_heading(f"Chunk {index}", level=2)

        self._add_key_value_paragraph(document, "Chunk ID", chunk.chunk_id)
        self._add_key_value_paragraph(document, "Strategy", chunk.strategy)
        self._add_key_value_paragraph(document, "Pages", f"{chunk.page_start} -> {chunk.page_end}")
        self._add_key_value_paragraph(
            document,
            "Source node type",
            getattr(chunk, "source_node_type", ""),
        )
        self._add_key_value_paragraph(
            document,
            "Source node label",
            getattr(chunk, "source_node_label", ""),
        )
        self._add_key_value_paragraph(
            document,
            "Chunk reason",
            getattr(chunk, "chunk_reason", ""),
        )
        self._add_key_value_paragraph(
            document,
            "Character count",
            str(getattr(chunk, "char_count", len(chunk.text or ""))),
        )
        self._add_key_value_paragraph(
            document,
            "Previous chunk",
            getattr(chunk, "prev_chunk_id", "") or "",
        )
        self._add_key_value_paragraph(
            document,
            "Next chunk",
            getattr(chunk, "next_chunk_id", "") or "",
        )

        hierarchy_path = list(getattr(chunk, "hierarchy_path", []) or [])
        if hierarchy_path:
            document.add_paragraph("Hierarchy path:")
            document.add_paragraph(
                sanitize_for_docx(" > ".join(hierarchy_path))
            )

        metadata = getattr(chunk, "metadata", {}) or {}
        if metadata:
            document.add_paragraph("Metadata:")
            for key in sorted(metadata.keys()):
                safe_key = sanitize_for_docx(str(key))
                safe_value = sanitize_for_docx(str(metadata[key]))
                document.add_paragraph(f"- {safe_key}: {safe_value}")

        document.add_paragraph("Visible text:")
        document.add_paragraph(sanitize_for_docx(chunk.text or ""))

        embedding_text = getattr(chunk, "text_for_embedding", "") or ""
        if embedding_text and embedding_text != (chunk.text or ""):
            document.add_paragraph("Embedding text:")
            document.add_paragraph(sanitize_for_docx(embedding_text))

    def _add_key_value_paragraph(
        self,
        document: Document,
        label: str,
        value: str,
    ) -> None:
        """
        Write one simple 'Label: Value' paragraph.

        Parameters
        ----------
        document : Document
            Target DOCX document object.

        label : str
            Field label.

        value : str
            Field value.
        """
        safe_label = sanitize_for_docx(label)
        safe_value = sanitize_for_docx(value or "")
        document.add_paragraph(f"{safe_label}: {safe_value}")